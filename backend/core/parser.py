"""
简历解析器 — 支持 PDF / Word / Markdown / TXT

原则：
1. 先用工具提取纯文本
2. 再用规则从纯文本中抽取结构化信息
3. 保留原始文本给 LLM 做深度分析
"""

import re
import os
from pathlib import Path
from typing import Optional, List, Dict, Tuple

# 懒加载 — 移动端调用不会装这些库的时候也能降级工作
_fitz = None
_docx = None

def _get_fitz():
    global _fitz
    if _fitz is None:
        try:
            import fitz as m
            _fitz = m
        except ImportError:
            pass
    return _fitz

def _get_docx():
    global _docx
    if _docx is None:
        try:
            import docx as m
            _docx = m
        except ImportError:
            pass
    return _docx


# ─── 文件读取 ────────────────────────────────────────────

def read_file_raw(filepath: str) -> Tuple[str, str]:
    """
    读取文件，返回 (raw_text, detected_format)
    支持: pdf, docx, doc, md, txt, rtf
    """
    path = Path(filepath)
    ext = path.suffix.lower()

    if ext == ".pdf":
        return _parse_pdf(path), "pdf"
    elif ext in (".docx", ".doc"):
        return _parse_docx(path), "word"
    elif ext in (".md", ".markdown"):
        return path.read_text(encoding="utf-8", errors="ignore"), "markdown"
    elif ext in (".txt", ".text", ""):
        return path.read_text(encoding="utf-8", errors="ignore"), "txt"
    elif ext == ".rtf":
        return _parse_rtf(path), "rtf"
    else:
        # 尝试按文本读
        try:
            return path.read_text(encoding="utf-8", errors="ignore"), "txt"
        except Exception:
            raise ValueError(f"不支持的文件格式: {ext}")


def _parse_pdf(path: Path) -> str:
    """用 pymupdf 提取 PDF 文本"""
    fitz = _get_fitz()
    if fitz is None:
        raise ImportError("需要安装 pymupdf: pip install pymupdf")

    doc = fitz.open(str(path))
    text_parts = []
    for page in doc:
        t = page.get_text()
        if t.strip():
            text_parts.append(t)
    doc.close()

    if not text_parts:
        # 尝试 OCR（扫描版 PDF）
        return _pdf_ocr(path)

    return "\n\n".join(text_parts)


def _pdf_ocr(path: Path) -> str:
    """扫描版 PDF 的 OCR 降级方案"""
    fitz = _get_fitz()
    doc = fitz.open(str(path))
    text_parts = []
    for i, page in enumerate(doc):
        # 提取嵌入图片做 OCR
        pix = page.get_pixmap(dpi=200)
        # 简化: 尝试用 page.get_text("dict") 提取文本块
        blocks = page.get_text("blocks")
        page_text = " ".join(b[4] for b in blocks if b[6] == 0 and b[4].strip())
        if page_text.strip():
            text_parts.append(page_text)
        else:
            text_parts.append(f"[第{i+1}页: 扫描图片，无法提取文本]")
    doc.close()
    return "\n\n".join(text_parts)


def _parse_docx(path: Path) -> str:
    """用 python-docx 提取 Word 文本"""
    docx = _get_docx()
    if docx is None:
        # 降级: 如果装了 libreoffice/pandoc 可以转
        raise ImportError("需要安装 python-docx: pip install python-docx")

    doc = docx.Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # 也提取表格内容
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    return "\n".join(paragraphs)


def _parse_rtf(path: Path) -> str:
    """简单的 RTF 文本提取（不依赖额外库）"""
    text = path.read_text(encoding="utf-8", errors="ignore")
    # 去掉 RTF 控制码 — 简化版
    text = re.sub(r'\\[a-zA-Z]+\d* ?', '', text)
    text = re.sub(r'[{}]', '', text)
    text = re.sub(r'\\\'[0-9a-fA-F]{2}', '', text)
    # 清理多余空白
    text = re.sub(r'\n\s*\n', '\n\n', text)
    return text.strip()


# ─── 结构抽取 ────────────────────────────────────────────

# 中文简历常见板块关键词
SECTION_PATTERNS = {
    "personal":   r"(个人(信息|资料|简介)|基本信息|联系方式|个人信息)",
    "summary":    r"(个人(简介|评价|总结)|自我(评价|描述|介绍)|求职意向|职业目标|关于我)",
    "skills":     r"(专业(技能|能力)|技术(栈|能力|擅长)|核心(技能|能力)|技能(特长|证书)?|掌握技能|编程语言|开发工具)",
    "work":       r"(工作(经历|经验|背景)|实习(经历|经验)|从业经历|职业经历|工作履历)",
    "project":    r"(项目(经历|经验|展示|案例)|主要项目|代表项目|项目作品)",
    "education":  r"(教育(?:背景|经历|信息)?|学历|学习经历|毕业院校)",
    "certificate": r"(证书|资格(证书)?|荣誉|获奖|奖项|资质|专业认证)",
    "language":   r"(语言(能力)?|外语|英语|日语)",
}

# 常见技能关键词库（按类别），用于从文本中提取
SKILL_KEYWORDS = {
    "hard_skill": [
        # 编程语言
        "Python", "Java", "JavaScript", "TypeScript", "Go", "Golang", "Rust",
        "C++", "C#", "PHP", "Ruby", "Swift", "Kotlin", "Dart", "Scala",
        "Shell", "Bash", "SQL", "MATLAB",
        # 前端
        "React", "Vue", "Angular", "Next.js", "Nuxt", "Svelte", "jQuery",
        "HTML", "CSS", "SCSS", "Less", "Tailwind", "Bootstrap", "Webpack", "Vite",
        "微信小程序", "uniapp", "Flutter", "React Native",
        # 后端
        "Django", "Flask", "FastAPI", "Spring", "Spring Boot", "MyBatis",
        "Node.js", "Express", "Koa", "NestJS", "Gin", "Beego",
        # 数据库
        "MySQL", "PostgreSQL", "MongoDB", "Redis", "Elasticsearch", "Oracle",
        "SQLite", "SQL Server", "DynamoDB", "ClickHouse", "TiDB",
        # 云/DevOps
        "Docker", "Kubernetes", "K8s", "Jenkins", "GitLab CI", "GitHub Actions",
        "AWS", "Azure", "阿里云", "腾讯云", "华为云", "GCP",
        "Terraform", "Ansible", "Nginx", "Linux",
        # 数据/AI
        "TensorFlow", "PyTorch", "Pandas", "NumPy", "Spark", "Hadoop",
        "Kafka", "RabbitMQ", "Flink", "机器学习", "深度学习", "NLP", "CV",
        # 设计
        "Figma", "Sketch", "Photoshop", "Illustrator", "After Effects",
        # 通用工具
        "Git", "SVN", "Jira", "Confluence",

        # ── IT 运维 / Helpdesk / 网络 ──
        "ITIL", "ITIL 4", "ITAM", "ITSM", "SLA", "OLA",
        "Windows", "MacOS", "macOS", "Linux", "Windows Server",
        "Active Directory", "AD 域", "域控", "AD域控",
        "Office 365", "Microsoft 365", "Exchange",
        "Helpdesk", "Help Desk", "桌面运维", "桌面支持",
        "IT 运维", "IT运维", "IT Support", "IT 支持",
        "事件管理", "问题管理", "变更管理", "资产管理", "配置管理",
        "CMDB", "IT资产", "IT 资产", "资产盘点",
        "TCP/IP", "DNS", "DHCP", "VPN", "LAN", "WAN", "VLAN",
        "网络配置", "网络排错", "网络运维", "网络故障",
        "打印机", "复印机", "投影仪", "会议系统",
        "Maxhub", "Cisco", "Polycom", "Zoom", "Teams",
        "Jamf", "Ivanti", "SCCM", "Intune", "MDM",
        "OBS", "推流", "直播推流",
        "OA", "OA系统", "OA流程", "泛微", "泛微OA", "致远", "钉钉", "飞书",
        "ServiceNow", "Jira Service Management", "Zendesk",
        "远程桌面", "TeamViewer", "AnyDesk", "VNC", "RDP",
        "备份", "灾备", "UPS", "机房", "IDC",
        "IT 资产管理", "IT资产管理", "硬件故障", "硬件排错",
        "电脑维修", "PC维修", "硬件维修",
        "服务台", "IT服务台", "Service Desk",

        # 产品/运营
        "数据分析", "用户研究", "A/B测试", "竞品分析", "需求文档", "PRD",
    ],
    "soft_skill": [
        "沟通能力", "团队协作", "项目管理", "领导力", "抗压能力",
        "逻辑思维", "问题解决", "创新", "学习能力", "执行力",
        "跨部门协作", "需求分析", "产品思维", "用户导向", "数据驱动",
        "敏捷开发", "Scrum", "OKR", "KPI",
    ],
    "certificate": [
        "PMP", "CFA", "CPA", "AWS认证", "RHCE", "CKA", "CKAD",
        "雅思", "托福", "CET-4", "CET-6", "专八", "N1", "N2",
        "软考", "建造师", "注册会计师",
    ],
}


def parse_to_resume(raw_text: str) -> dict:
    """
    从纯文本中抽取结构化简历信息
    返回 dict，可传给 Resume(**dict)
    """
    # 1. 分块
    sections = _split_sections(raw_text)

    # 2. 逐块解析
    result = {
        "name": "",
        "email": "",
        "phone": "",
        "city": "",
        "title": "",
        "summary": "",
        "hard_skills": [],
        "soft_skills": [],
        "languages": [],
        "certificates": [],
        "work_experience": [],
        "projects": [],
        "education": [],
        "years_of_experience": 0.0,
        "raw_text": raw_text,
    }

    # 提取联系信息
    result["email"] = _extract_email(raw_text)
    result["phone"] = _extract_phone(raw_text)

    # 解析各板块
    for section_type, content in sections.items():
        if section_type == "personal":
            # 名字、城市等
            result["name"] = _extract_name(content, raw_text)
            result["city"] = _extract_city(raw_text)

        elif section_type == "summary":
            result["summary"] = content
            result["title"] = _extract_title(content)

        elif section_type == "skills":
            result["hard_skills"], result["soft_skills"] = _extract_skills(content)
            # 也从全文中提取技能（有些简历把技能散落在经历中）
            h2, s2 = _extract_skills(raw_text)
            result["hard_skills"] = list(set(result["hard_skills"] + h2))
            result["soft_skills"] = list(set(result["soft_skills"] + s2))

        elif section_type == "work":
            result["work_experience"] = _extract_sections(content, "work")

        elif section_type == "project":
            result["projects"] = _extract_sections(content, "project")

        elif section_type == "education":
            result["education"] = _extract_sections(content, "education")

        elif section_type == "certificate":
            result["certificates"] = _extract_certificates(content)

        elif section_type == "language":
            result["languages"] = _extract_languages(content)

    # 估算工作年限
    result["years_of_experience"] = _estimate_years(result["work_experience"])

    return result


def _split_sections(text: str) -> Dict[str, str]:
    """将简历文本按板块切分"""
    # 找到所有板块标题的位置
    positions = []
    for section_name, pattern in SECTION_PATTERNS.items():
        for m in re.finditer(pattern, text, re.IGNORECASE):
            positions.append((m.start(), section_name))

    if not positions:
        return {"_other": text}

    positions.sort()
    sections = {}

    for i, (pos, name) in enumerate(positions):
        start = pos
        if i + 1 < len(positions):
            end = positions[i + 1][0]
        else:
            end = len(text)
        sections[name] = text[start:end].strip()

    # 第一板块之前的内容
    if positions[0][0] > 0:
        prefix = text[:positions[0][0]].strip()
        if prefix:
            sections["personal"] = prefix

    return sections


def _extract_email(text: str) -> str:
    m = re.search(r'[\w.+-]+@[\w-]+\.[\w.-]+', text)
    return m.group(0) if m else ""


def _extract_phone(text: str) -> str:
    # 中国大陆手机号
    m = re.search(r'1[3-9]\d{9}', text)
    return m.group(0) if m else ""


def _extract_name(text: str, full_text: str = "") -> str:
    """提取姓名，通常在最开头"""
    lines = [l.strip() for l in (text or full_text).split("\n") if l.strip()]
    for line in lines[:5]:
        # 如果行中有 | / 分隔符，取第一部分作为候选名
        parts = re.split(r'\s*[|/]\s*', line)
        first_part = parts[0].strip()

        # 去掉邮箱、电话
        first_part = re.sub(r'[\w.+-]+@[\w-]+\.[\w.-]+', '', first_part)
        first_part = re.sub(r'1[3-9]\d{9}', '', first_part)
        first_part = first_part.strip()

        # 中文名 2-4 字
        m = re.fullmatch(r'[\u4e00-\u9fa5]{2,4}', first_part)
        if m:
            return m.group(0)

        # 检查整个 cleaned line
        cleaned = re.sub(r'[\w.+-]+@[\w-]+\.[\w.-]+', '', line)
        cleaned = re.sub(r'1[3-9]\d{9}', '', cleaned)
        cleaned = cleaned.strip().strip("|").strip()

        m = re.search(r'[\u4e00-\u9fa5]{2,4}', cleaned)
        if m and len(cleaned) <= 10:
            return m.group(0)

        # 英文名
        m = re.search(r'[A-Z][a-z]+ [A-Z][a-z]+', cleaned)
        if m:
            return m.group(0)
    return ""


def _extract_city(text: str) -> str:
    """提取所在城市"""
    cities = [
        "北京", "上海", "广州", "深圳", "杭州", "成都", "南京", "武汉",
        "西安", "重庆", "苏州", "天津", "长沙", "郑州", "东莞", "青岛",
        "沈阳", "宁波", "昆明", "大连", "厦门", "合肥", "佛山", "福州",
        "哈尔滨", "济南", "温州", "长春", "石家庄", "常州", "泉州", "南宁",
        "贵阳", "南昌", "太原", "烟台", "嘉兴", "南通", "金华", "珠海",
        "惠州", "徐州", "海口", "乌鲁木齐", "兰州", "呼和浩特", "银川",
        "西宁", "拉萨", "香港", "澳门", "台北",
    ]
    for city in cities:
        if city in text:
            return city
    return ""


def _extract_title(text: str) -> str:
    """从摘要中提取期望职位"""
    titles = re.findall(r'(?:期望职位|求职意向|目标岗位|意向岗位)[：:]\s*(.+?)(?:[，,。\n]|$)', text)
    if titles:
        return titles[0].strip()

    # 常见职位名匹配
    common_titles = [
        "前端", "后端", "全栈", "架构师", "技术总监", "CTO",
        "产品经理", "项目经理", "UI设计师", "UX设计师", "交互设计",
        "数据分析师", "算法工程师", "测试工程师", "运维工程师",
        "运营", "市场", "销售", "财务", "HR", "行政",
        "Java", "Python", "Go", "C++",
    ]
    for title in common_titles:
        if title in text:
            # 找完整职位名
            m = re.search(rf'[\u4e00-\u9fa5a-zA-Z+]*{title}[\u4e00-\u9fa5a-zA-Z]*', text)
            if m:
                # 检查是否是职位描述的一部分
                ctx = m.group(0)
                if any(kw in text[max(0,m.start()-20):m.end()+20]
                       for kw in ["工程师", "经理", "主管", "专员", "总监", "负责人", "开发"]):
                    return ctx
    return ""


def _extract_skills(text: str) -> Tuple[List[str], List[str]]:
    """提取硬技能和软技能"""
    hard = []
    soft = []
    text_lower = text.lower()

    for skill in SKILL_KEYWORDS["hard_skill"]:
        if skill.lower() in text_lower:
            hard.append(skill)

    for skill in SKILL_KEYWORDS["soft_skill"]:
        if skill in text:
            soft.append(skill)

    return hard, soft


def _extract_certificates(text: str) -> List[str]:
    """提取证书"""
    found = []
    for cert in SKILL_KEYWORDS["certificate"]:
        if cert.lower() in text.lower():
            found.append(cert)
    return found


def _extract_languages(text: str) -> List[str]:
    """提取语言能力"""
    langs = []
    lang_map = {
        "英语": ["英语", "English", "CET-4", "CET-6", "雅思", "托福", "专八", "专四"],
        "日语": ["日语", "Japanese", "N1", "N2", "N3", "N4", "N5"],
        "韩语": ["韩语", "Korean"],
        "法语": ["法语", "French"],
        "德语": ["德语", "German"],
        "西班牙语": ["西班牙语", "Spanish"],
    }
    for lang_name, keywords in lang_map.items():
        for kw in keywords:
            if kw.lower() in text.lower():
                langs.append(lang_name)
                break
    return langs


def _extract_sections(text: str, section_type: str) -> List[dict]:
    """
    从经历版块中拆分各段经历
    返回 List[dict] 结构
    """
    # 去掉版块标题行（如 "工作经历:" "项目经历:"）
    lines = text.strip().split("\n")
    if lines:
        first = re.sub(r'【.*?】', '', lines[0]).strip().rstrip("：:")
        if any(kw in first for kw in ["工作经历", "工作经验", "项目经历", "教育背景",
                                        "教育经历", "个人技能", "专业技能", "证书",
                                        "自我评价", "个人简介", "语言能力"]):
            text = "\n".join(lines[1:]).strip()

    if not text.strip():
        return []

    entries = []

    # 策略：先用空行拆成大块，块内再解析公司/日期/描述
    blocks = re.split(r'\n\s*\n', text)
    
    for block in blocks:
        block = block.strip()
        if len(block) < 10:
            continue
            
        # 尝试解析这个 block
        lines = block.split("\n")
        
        # 如果 block 的第一行是日期行，把它和之前的内容合并
        # (处理 "公司名\n日期\n描述" 格式)
        date_header = re.compile(r'^(\d{4}(?:[./-]\d{1,2})?)\s*[-–—至到]\s*(\d{4}(?:[./-]\d{1,2})?|至今|现在|present)')
        
        # 把 block 内的行按日期行重新分组
        sub_blocks = []
        current_lines = []
        
        for line in lines:
            stripped = line.strip()
            if date_header.match(stripped) and current_lines:
                sub_blocks.append(current_lines)
                current_lines = [line]
            else:
                current_lines.append(line)
        if current_lines:
            sub_blocks.append(current_lines)
        
        # 如果第一个 sub_block 只有 1-2 行且不含日期，它可能是公司抬头
        # 把它和下一个含日期的 sub_block 合并
        merged_blocks = []
        i = 0
        while i < len(sub_blocks):
            sb = sub_blocks[i]
            has_date = any(date_header.match(l.strip()) for l in sb)
            
            if not has_date and len(sb) <= 2 and i + 1 < len(sub_blocks):
                # 合并到下一个 block
                merged = sb + sub_blocks[i + 1]
                merged_blocks.append(merged)
                i += 2
            else:
                merged_blocks.append(sb)
                i += 1
        
        for sb in merged_blocks:
            chunk_text = "\n".join(sb).strip()
            if len(chunk_text) < 10:
                continue
            entry = _parse_entry(chunk_text, section_type)
            if entry["organization"] or entry["description"]:
                entries.append(entry)

    return entries


def _parse_entry(text: str, section_type: str) -> dict:
    """解析单段经历"""
    entry = {
        "title": "",
        "organization": "",
        "role": "",
        "start_date": "",
        "end_date": "",
        "description": "",
        "highlights": [],
        "skills_used": [],
    }

    lines = [l.strip() for l in text.split("\n") if l.strip()]

    # 第一行通常是 公司/学校 + 职位/专业
    if lines:
        first = lines[0]
        # 提取时间 — 先看第一行，没有再看后续行
        time_m = None
        for line_idx, line in enumerate(lines[:3]):  # 检查前三行
            time_m = re.search(r'(\d{4}(?:[./-]\d{1,2})?)\s*[-–—至到]\s*(\d{4}(?:[./-]\d{1,2})?|至今|现在|present)',
                              line.strip())
            if time_m:
                entry["start_date"] = time_m.group(1).replace("/", "-").replace(".", "-")
                entry["end_date"] = time_m.group(2).replace("/", "-").replace(".", "-")
                if line_idx == 0:
                    # 日期在第一行 — 取日期前后的文本
                    before = line[:time_m.start()].strip().rstrip("|/·• ").strip()
                    after = line[time_m.end():].strip().lstrip("|/·• ").strip()
                    cleaned = before if before else after
                else:
                    # 日期在后续行，第一行是组织
                    cleaned = first.strip()
                break
        
        if not time_m:
            cleaned = first.strip()

        # 尝试按 | / · 分割
        parts = re.split(r'\s*[|/·•]\s*', cleaned)
        parts = [p.strip() for p in parts if p.strip()]

        if section_type == "education":
            if len(parts) >= 2:
                entry["organization"] = parts[0]
                entry["role"] = parts[1]  # 学历/专业
            else:
                entry["organization"] = cleaned
        else:
            if len(parts) >= 2:
                entry["organization"] = parts[0]
                entry["role"] = parts[1]
            elif len(parts) == 1:
                # 可能公司名包含了职位
                entry["organization"] = cleaned
            entry["title"] = cleaned

    # 剩余行是描述
    if len(lines) > 1:
        entry["description"] = "\n".join(lines[1:])

        # 提取带 · - * 开头的亮点
        for line in lines[1:]:
            m = re.match(r'^[·•\-\*\d\.]\s*(.+)', line)
            if m:
                entry["highlights"].append(m.group(1).strip())

    # 提取描述中用到的技能
    hard, _ = _extract_skills(entry["description"])
    entry["skills_used"] = hard

    return entry


def _estimate_years(work_experience: List[dict]) -> float:
    """估算工作年限"""
    if not work_experience:
        return 0.0

    # 收集所有年份
    years_set = set()
    for exp in work_experience:
        sd = exp.get("start_date", "")
        ed = exp.get("end_date", "")

        # start_date 可能格式: "2020", "2020-03", "2020.03"
        if sd:
            sd_clean = sd.replace(".", "-")
            m = re.match(r'(\d{4})', sd_clean)
            if m:
                year = int(m.group(1))
                years_set.add(year)

        if ed and ed not in ("至今", "现在", "present", ""):
            ed_clean = ed.replace(".", "-")
            m = re.match(r'(\d{4})', ed_clean)
            if m:
                year = int(m.group(1))
                years_set.add(year)
        elif ed in ("至今", "现在", "present"):
            # 至今 → 算到当前年份
            from datetime import datetime
            years_set.add(datetime.now().year)

    if len(years_set) >= 2:
        return float(max(years_set) - min(years_set))
    elif len(years_set) == 1:
        # 只有一年，无法估算
        return 1.0

    return 0.0


# ─── 主入口 ─────────────────────────────────────────────

def parse_resume_file(filepath: str) -> dict:
    """
    解析简历文件，返回结构化 dict

    Usage:
        result = parse_resume_file("/path/to/resume.pdf")
        resume = Resume(**result)
    """
    raw_text, fmt = read_file_raw(filepath)
    result = parse_to_resume(raw_text)
    result["_format"] = fmt
    result["_raw_length"] = len(raw_text)
    return result
